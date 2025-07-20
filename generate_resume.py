# generate_resume.py

"""
AI-Powered Resume Generation Module for JobFinder

This module generates personalized resumes by:
1. Analyzing job descriptions for key requirements and keywords
2. Tailoring the user's base resume to match job requirements
3. Optimizing for ATS (Applicant Tracking Systems)
4. Creating professionally formatted Word documents
"""

import re
import json
from pathlib import Path
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import openai
from config import load

# For debugging resume generation
import sys

def clean_company_name(company: str) -> str:
    """Clean and deduplicate company name"""
    # Clean up common suffixes and extra spaces
    company = re.sub(r'\s+(Inc|LLC|Corp|Ltd|Co)\.?$', '', company, flags=re.IGNORECASE)
    company = re.sub(r'\s+', ' ', company)  # Replace multiple spaces with single space
    
    # Remove duplicate words within the company name
    words = company.split()
    unique_words = []
    seen_words = set()
    for word in words:
        word_lower = word.lower()
        if word and word_lower not in seen_words:
            unique_words.append(word)
            seen_words.add(word_lower)
    company = ' '.join(unique_words)
    
    # Additional check for exact string repetitions (like "SilverchairSilverchair")
    # Look for patterns where the same substring is repeated consecutively
    original_company = company
    if len(company) > 6:  # Only check longer strings
        # Check if the company name is a repetition of itself
        for i in range(2, len(company) // 2 + 1):  # Start from 2 chars minimum
            prefix = company[:i]
            if len(prefix) > 1 and company == prefix + prefix:  # Simple case: exact duplication
                company = prefix
                print(f"🏢 DEBUG: Detected exact duplication, reduced '{original_company}' to '{company}'")
                break
            elif company.startswith(prefix) and company.endswith(prefix) and len(company) == len(prefix) * 2:
                company = prefix
                print(f"🏢 DEBUG: Detected prefix/suffix repetition, reduced '{original_company}' to '{company}'")
                break
    
    return company

def extract_company_name(job_title: str, job_description: str) -> str:
    """Extract company name from job title or description"""
    
    print(f"🏢 DEBUG: Input job_title: '{job_title}'")
    print(f"🏢 DEBUG: Input job_description (first 200 chars): '{job_description[:200]}...'")
    
    # Common patterns in job titles
    title_patterns = [
        r'at\s+([A-Z][a-zA-Z\s&.,]+?)(?:\s*[-|]|$)',
        r'@\s*([A-Z][a-zA-Z\s&.,]+?)(?:\s*[-|]|$)',
        r'-\s*([A-Z][a-zA-Z\s&.,]+?)(?:\s*[-|]|$)',
    ]
    
    for pattern in title_patterns:
        match = re.search(pattern, job_title, re.IGNORECASE)
        if match:
            company = match.group(1).strip()
            print(f"🏢 DEBUG: Found company in title with pattern '{pattern}': '{company}'")
            company = clean_company_name(company)
            print(f"🏢 DEBUG: After cleaning: '{company}'")
            if len(company) > 2 and len(company) < 50:
                return company
    
    # Patterns in job description
    description_patterns = [
        r'(?:About|Join|At)\s+([A-Z][a-zA-Z\s&.,]+?)(?:\s*[,\n]|\s+is\s|\s+are\s)',
        r'([A-Z][a-zA-Z\s&.,]+?)\s+is\s+(?:a\s+|an\s+)?(?:leading|growing|innovative|established)',
        r'Company:\s*([A-Z][a-zA-Z\s&.,]+?)(?:\s*\n|$)',
        r'Organization:\s*([A-Z][a-zA-Z\s&.,]+?)(?:\s*\n|$)',
    ]
    
    for pattern in description_patterns:
        match = re.search(pattern, job_description[:1000], re.IGNORECASE)  # Search first 1000 chars
        if match:
            company = match.group(1).strip()
            print(f"🏢 DEBUG: Found company in description with pattern '{pattern}': '{company}'")
            company = clean_company_name(company)
            print(f"🏢 DEBUG: After cleaning: '{company}'")
            if len(company) > 2 and len(company) < 50:
                return company
    
    print(f"🏢 DEBUG: No company name found, using fallback")
    return "Company"  # Fallback

def generate_personalized_resume_content(original_resume: str, job_description: str, job_title: str) -> Dict[str, str]:
    """Use AI to generate personalized resume content based on job requirements"""
    
    # Load fresh config each time (exactly like evaluate.py does)
    current_config = load()
    openai_api_key = current_config.get("api_keys", {}).get("openai_api_key")
    
    # Clean API key of any whitespace/newlines (defensive programming)
    if openai_api_key:
        openai_api_key = openai_api_key.strip().replace('\n', '').replace('\r', '').replace(' ', '')
    
    # Debug: Show what we loaded (truncate key for security)
    key_display = f"{openai_api_key[:15]}...{openai_api_key[-10:]}" if openai_api_key and len(openai_api_key) > 25 else openai_api_key
    print(f"🔧 DEBUG: Cleaned API key: {key_display} (length: {len(openai_api_key) if openai_api_key else 0})")
    
    # Validate API key (matching evaluate.py pattern exactly)
    if not openai_api_key or openai_api_key == "YOUR_OPENAI_API_KEY_HERE":
        raise ValueError("OpenAI API Key not configured in config.toml or is a placeholder.")
    
    # Set up OpenAI client (matching evaluate.py pattern exactly)
    openai.api_key = openai_api_key
    
    enhanced_prompt = f"""
You are an expert resume writer and ATS optimization specialist. Your task is to create a comprehensive, personalized resume for a specific job posting by analyzing the candidate's background and the job requirements.

CANDIDATE'S ORIGINAL RESUME:
{original_resume}

TARGET JOB TITLE: {job_title}

TARGET JOB DESCRIPTION:
{job_description}

DETAILED INSTRUCTIONS:
1. ANALYZE the job description to identify:
   - Required technical skills and technologies
   - Preferred qualifications and experience levels  
   - Key responsibilities and job functions
   - Industry-specific keywords and terminology
   - Company culture and values mentioned

2. TRANSFORM the original resume by:
   - Extracting ALL relevant work experience, education, and skills from the original resume
   - Rewriting job descriptions to emphasize experiences relevant to this target role
   - Incorporating job-specific keywords naturally throughout the content
   - Quantifying achievements with numbers, percentages, and metrics where possible
   - Using strong action verbs that match the job posting language

3. OPTIMIZE for ATS systems by:
   - Including exact keyword matches from the job description
   - Using standard section headings and formatting
   - Avoiding graphics, tables, or unusual formatting
   - Ensuring keyword density without stuffing

4. CREATE comprehensive sections with substantial content:
   - Professional Summary: 3-4 compelling sentences that position the candidate perfectly for this role
   - Key Skills: 10-15 relevant technical and soft skills from both the original resume and job requirements
   - Work Experience: ALL relevant positions with 3-5 detailed, tailored achievement bullets each
   - Education: Complete educational background
   - Additional relevant sections (certifications, projects, etc.)

Return ONLY a valid JSON object with this exact structure:
{{
    "professional_summary": "A compelling 3-4 sentence summary that positions the candidate perfectly for this specific role, incorporating key job requirements and the candidate's strongest qualifications",
    "key_skills": ["skill1", "skill2", "skill3", "skill4", "skill5", "skill6", "skill7", "skill8", "skill9", "skill10", "skill11", "skill12"],
    "work_experience": [
        {{
            "title": "Exact Job Title from Original Resume",
            "company": "Company Name", 
            "duration": "Month Year - Month Year",
            "achievements": [
                "Detailed achievement using action verb and quantified results relevant to target role",
                "Second achievement emphasizing skills mentioned in job posting", 
                "Third achievement showing impact and relevant experience",
                "Fourth achievement if applicable"
            ]
        }}
    ],
    "education": [
        {{
            "degree": "Full Degree Name",
            "institution": "University/School Name",
            "year": "Graduation Year or Years Attended"
        }}
    ],
    "additional_sections": {{
        "certifications": ["Certification 1", "Certification 2"],
        "projects": ["Relevant project 1", "Relevant project 2"],
        "languages": ["Language 1", "Language 2"],
        "technical_proficiencies": ["Technology 1", "Technology 2"]
    }}
}}

CRITICAL: 
- Extract and include ALL relevant work experience from the original resume
- Create detailed, specific achievement bullets for each position
- Ensure the final resume would be substantial and comprehensive, not minimal
- Keep all factual information accurate while optimizing presentation
- The resume should be rich in content and perfectly tailored to this specific job
"""

    try:
        print(f"Generating personalized resume using OpenAI (o3 → gpt-4o fallback) for job: {job_title}")
        
        # Try o3 first (like evaluate.py), then fallback options
        models_to_try = ["o3", "gpt-4o", "gpt-4o-mini"]
        response = None
        
        for model_name in models_to_try:
            try:
                print(f"Attempting to use model: {model_name}")
                
                # Determine parameters based on model type
                if model_name.startswith("o3"):
                    # o3 models have strict parameter requirements
                    token_param = "max_completion_tokens"
                    token_limit = 25000  # High limit to account for internal reasoning tokens
                    print(f"🧠 Using o3-optimized parameters: {token_param}={token_limit}, temperature=default")
                    
                    # Build o3-specific request parameters (no custom temperature allowed)
                    request_params = {
                        "model": model_name,
                        "messages": [
                            {
                                "role": "system", 
                                "content": "You are an expert resume writer with 15+ years of experience in ATS optimization and job-specific personalization. You create comprehensive, detailed resumes that perfectly match job requirements while showcasing all relevant candidate qualifications."
                            },
                            {"role": "user", "content": enhanced_prompt}
                        ],
                        token_param: token_limit
                        # NOTE: No temperature parameter - o3 only supports default (1)
                    }
                else:
                    # Legacy models use max_tokens with standard parameters
                    token_param = "max_tokens" 
                    token_limit = 3000
                    print(f"🔧 Using legacy parameters: {token_param}={token_limit}, temperature=0.2")
                    
                    # Build legacy request parameters
                    request_params = {
                        "model": model_name,
                        "messages": [
                            {
                                "role": "system", 
                                "content": "You are an expert resume writer with 15+ years of experience in ATS optimization and job-specific personalization. You create comprehensive, detailed resumes that perfectly match job requirements while showcasing all relevant candidate qualifications."
                            },
                            {"role": "user", "content": enhanced_prompt}
                        ],
                        "temperature": 0.2,
                        token_param: token_limit
                    }
                
                # Use OpenAI client with dynamic parameters
                response = openai.chat.completions.create(**request_params)
                
                print(f"Successfully used model: {model_name}")
                break  # Success, exit the loop
            except Exception as model_error:
                error_msg = str(model_error)
                print(f"Model {model_name} failed: {error_msg}")
                
                # Check if it's an API key issue (match evaluate.py error handling)
                if "invalid_api_key" in error_msg or "401" in error_msg:
                    print("❌ API key error detected!")
                    raise ValueError("OpenAI API Key not configured in config.toml or is a placeholder.")
                
                if model_name == models_to_try[-1]:  # Last model in list
                    raise  # Re-raise the error if all models fail
                continue  # Try next model
        
        if not response:
            raise Exception("All AI models failed to generate response")
        
        # Extract content (same for both client versions)
        content = response.choices[0].message.content.strip()
        print(f"AI Response received, length: {len(content)} characters")
        
        # Clean the response to ensure it's valid JSON
        if content.startswith("```json"):
            content = content[7:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        # Parse JSON response
        try:
            resume_data = json.loads(content)
            print(f"Resume data parsed successfully. Sections: {list(resume_data.keys())}")
            
            # Validate that we have substantial content
            work_exp_count = len(resume_data.get("work_experience", []))
            skills_count = len(resume_data.get("key_skills", []))
            
            if work_exp_count == 0:
                print("Warning: No work experience found in AI response")
            else:
                print(f"AI generated {work_exp_count} work experience entries")
                
            if skills_count == 0:
                print("Warning: No skills found in AI response")
            else:
                print(f"AI generated {skills_count} skills")
            
            return resume_data
            
        except json.JSONDecodeError as je:
            print(f"JSON parsing error: {je}")
            print(f"Raw AI response: {content[:500]}...")
            
            # Try to extract information from non-JSON response
            return parse_non_json_response(content, original_resume)
            
    except Exception as e:
        print(f"Error generating resume content with AI: {e}")
        
        # Enhanced fallback - try to extract basic info from original resume
        return create_fallback_resume(original_resume, job_title, job_description)

def parse_non_json_response(content: str, original_resume: str) -> Dict[str, str]:
    """Parse non-JSON AI response and extract resume information"""
    print("Attempting to parse non-JSON response...")
    
    # This is a fallback parser for when AI doesn't return proper JSON
    return create_fallback_resume(original_resume, "", "")

def create_fallback_resume(original_resume: str, job_title: str, job_description: str) -> Dict[str, str]:
    """Create a basic resume structure when AI fails"""
    print("Creating fallback resume from original content...")
    
    # Try to extract basic information from original resume
    lines = [line.strip() for line in original_resume.split('\n') if line.strip()]
    
    # Extract skills
    skills = []
    for line in lines:
        if any(keyword in line.lower() for keyword in ['skill', 'technology', 'programming', 'software']):
            # Extract technology names
            tech_keywords = ['python', 'javascript', 'java', 'react', 'node', 'sql', 'html', 'css', 'git']
            for tech in tech_keywords:
                if tech in line.lower():
                    skills.append(tech.title())
    
    if not skills:
        skills = ["Communication", "Problem Solving", "Team Collaboration", "Project Management"]
    
    # Try to find work experience
    work_experience = []
    current_job = None
    
    for line in lines:
        # Look for job titles/companies
        if any(word in line.lower() for word in ['engineer', 'developer', 'manager', 'analyst', 'coordinator']):
            if current_job:
                work_experience.append(current_job)
            current_job = {
                "title": line,
                "company": "Previous Company",
                "duration": "2020 - Present",
                "achievements": [
                    "Contributed to team objectives and project deliverables",
                    "Collaborated with cross-functional teams",
                    "Applied technical skills to solve complex problems"
                ]
            }
    
    if current_job:
        work_experience.append(current_job)
    
    return {
        "professional_summary": f"Experienced professional with relevant background seeking to contribute skills and expertise to {job_title or 'this role'}. Strong track record of delivering results and working effectively in team environments.",
        "key_skills": skills[:12],
        "work_experience": work_experience,
        "education": [
            {
                "degree": "Relevant Degree",
                "institution": "University",
                "year": "Year"
            }
        ],
        "additional_sections": {
            "certifications": [],
            "projects": [],
            "languages": []
        }
    }

def create_professional_resume_document(
    resume_data: Dict[str, Any], 
    personal_info: Dict[str, str], 
    company_name: str
) -> Document:
    """Create a professionally formatted Word document"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Create custom styles
    styles = doc.styles
    
    # Header style
    if 'HeaderStyle' not in [s.name for s in styles]:
        header_style = styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
        header_font = header_style.font
        header_font.name = 'Calibri'
        header_font.size = Pt(16)
        header_font.bold = True
        header_font.color.rgb = RGBColor(0x2F, 0x54, 0x96)  # Professional blue
        header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_style.paragraph_format.space_after = Pt(6)
    
    # Contact info style
    if 'ContactStyle' not in [s.name for s in styles]:
        contact_style = styles.add_style('ContactStyle', WD_STYLE_TYPE.PARAGRAPH)
        contact_font = contact_style.font
        contact_font.name = 'Calibri'
        contact_font.size = Pt(10)
        contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_style.paragraph_format.space_after = Pt(12)
    
    # Section heading style
    if 'SectionHeading' not in [s.name for s in styles]:
        section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        section_font = section_style.font
        section_font.name = 'Calibri'
        section_font.size = Pt(12)
        section_font.bold = True
        section_font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        section_style.paragraph_format.space_before = Pt(8)
        section_style.paragraph_format.space_after = Pt(4)
        # Add bottom border
        pPr = section_style.element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2F5496')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    # Body text style
    if 'BodyText' not in [s.name for s in styles]:
        body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(10)
        body_style.paragraph_format.space_after = Pt(3)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Add header with name and contact info
    name_para = doc.add_paragraph(f"{personal_info.get('first_name', '')} {personal_info.get('last_name', '')}".strip(), style='HeaderStyle')
    
    # Contact information
    contact_info = []
    if personal_info.get('phone_number'):
        contact_info.append(personal_info['phone_number'])
    if personal_info.get('email_address'):
        contact_info.append(personal_info['email_address'])
    if personal_info.get('linkedin_url'):
        linkedin_display = personal_info['linkedin_url'].replace('https://', '').replace('http://', '')
        contact_info.append(linkedin_display)
    
    if contact_info:
        contact_para = doc.add_paragraph(' | '.join(contact_info), style='ContactStyle')
    
    # Professional Summary
    if resume_data.get('professional_summary'):
        doc.add_paragraph('PROFESSIONAL SUMMARY', style='SectionHeading')
        doc.add_paragraph(resume_data['professional_summary'], style='BodyText')
    
    # Key Skills
    if resume_data.get('key_skills'):
        doc.add_paragraph('KEY SKILLS', style='SectionHeading')
        skills_text = ' • '.join(resume_data['key_skills'])
        doc.add_paragraph(skills_text, style='BodyText')
    
    # Work Experience
    if resume_data.get('work_experience'):
        doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='SectionHeading')
        for exp in resume_data['work_experience']:
            # Job title and company
            job_para = doc.add_paragraph(style='BodyText')
            job_run = job_para.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}")
            job_run.bold = True
            
            # Duration
            if exp.get('duration'):
                duration_para = doc.add_paragraph(exp['duration'], style='BodyText')
                duration_para.paragraph_format.space_after = Pt(2)
            
            # Achievements
            if exp.get('achievements'):
                for achievement in exp['achievements']:
                    bullet_para = doc.add_paragraph(f"• {achievement}", style='BodyText')
                    bullet_para.paragraph_format.left_indent = Inches(0.2)
    
    # Education
    if resume_data.get('education'):
        doc.add_paragraph('EDUCATION', style='SectionHeading')
        for edu in resume_data['education']:
            edu_text = f"{edu.get('degree', '')} | {edu.get('institution', '')}"
            if edu.get('year'):
                edu_text += f" | {edu['year']}"
            doc.add_paragraph(edu_text, style='BodyText')
    
    # Additional sections
    additional = resume_data.get('additional_sections', {})
    if additional.get('certifications'):
        doc.add_paragraph('CERTIFICATIONS', style='SectionHeading')
        cert_text = ' • '.join(additional['certifications'])
        doc.add_paragraph(cert_text, style='BodyText')
    
    if additional.get('projects'):
        doc.add_paragraph('NOTABLE PROJECTS', style='SectionHeading')
        for project in additional['projects'][:3]:  # Limit to 3 for space
            doc.add_paragraph(f"• {project}", style='BodyText')
    
    return doc

def generate_resume(
    job_title: str,
    job_description: str, 
    original_resume: str,
    personal_info: Dict[str, str]
) -> tuple[bytes, str]:
    """
    Main function to generate a personalized resume
    
    Returns:
        tuple: (document_bytes, filename)
    """
    
    # Extract company name for filename
    company_name = extract_company_name(job_title, job_description)
    
    # Generate AI-optimized resume content
    resume_data = generate_personalized_resume_content(
        original_resume, 
        job_description, 
        job_title
    )
    
    # Create Word document
    doc = create_professional_resume_document(
        resume_data, 
        personal_info, 
        company_name
    )
    
    # Generate filename
    first_name = personal_info.get('first_name', 'Resume').replace(' ', '')
    last_name = personal_info.get('last_name', '').replace(' ', '')
    
    # Clean company name for filename (company_name should already be cleaned by extract_company_name)
    company_clean = re.sub(r'[^\w\s-]', '', company_name)
    company_clean = re.sub(r'\s+', '_', company_clean.strip())  # Replace spaces with underscores
    
    filename = f"{first_name}_{last_name}_{company_clean}.docx"
    print(f"📄 Generated filename: {filename}")
    
    # Convert document to bytes
    from io import BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_bytes = doc_buffer.getvalue()
    doc_buffer.close()
    
    return doc_bytes, filename

# Test function for development
def test_resume_generation():
    """Test function to verify resume generation works"""
    sample_personal_info = {
        "first_name": "John",
        "last_name": "Doe", 
        "phone_number": "(555) 123-4567",
        "email_address": "john.doe@email.com",
        "linkedin_url": "https://linkedin.com/in/johndoe"
    }
    
    sample_resume = """
    John Doe - Software Engineer
    5+ years of experience in full-stack web development
    
    Professional Experience:
    Software Developer - Tech Corp (January 2019 - Present)
    • Developed and maintained web applications using React, Node.js, and Python
    • Collaborated with cross-functional teams of 8+ developers and designers
    • Improved application performance by 40% through code optimization
    • Mentored 3 junior developers and conducted code reviews
    
    Junior Developer - StartupCo (June 2018 - December 2018)
    • Built responsive web interfaces using HTML, CSS, and JavaScript
    • Worked with RESTful APIs and integrated third-party services
    • Participated in agile development process and daily standups
    
    Technical Skills:
    • Programming Languages: Python, JavaScript, TypeScript, HTML, CSS
    • Frameworks: React, Node.js, Express, Django, Flask
    • Databases: MySQL, PostgreSQL, MongoDB
    • Tools: Git, Docker, AWS, Jenkins
    
    Education:
    Bachelor of Science in Computer Science - State University (2018)
    Relevant coursework: Data Structures, Algorithms, Software Engineering
    
    Certifications:
    • AWS Certified Developer Associate (2022)
    • React Developer Certification (2021)
    """
    
    sample_job_description = """
    Python Developer - Innovation Labs
    
    About Innovation Labs:
    Innovation Labs is a fast-growing technology company specializing in AI-powered solutions.
    
    Job Description:
    We are seeking a talented Python Developer to join our engineering team. The ideal candidate will have strong experience in Python development and web frameworks.
    
    Requirements:
    • 3+ years of Python programming experience
    • Experience with Django or Flask web frameworks
    • Strong knowledge of SQL and database design
    • Familiarity with RESTful API development
    • Experience with version control (Git)
    • Strong problem-solving and analytical skills
    • Bachelor's degree in Computer Science or related field
    
    Preferred Qualifications:
    • Experience with cloud platforms (AWS, Azure, GCP)
    • Knowledge of containerization (Docker)
    • Experience with automated testing
    • Understanding of agile development methodologies
    
    What We Offer:
    • Competitive salary and benefits
    • Remote work flexibility
    • Professional development opportunities
    • Collaborative team environment
    """
    
    try:
        print("🧪 Testing resume generation...")
        doc_bytes, filename = generate_resume(
            "Python Developer at Innovation Labs",
            sample_job_description,
            sample_resume,
            sample_personal_info
        )
        print(f"✅ Resume generated successfully!")
        print(f"📄 Filename: {filename}")
        print(f"📊 Document size: {len(doc_bytes)} bytes")
        return True
    except Exception as e:
        print(f"❌ Resume generation failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_company_name_extraction():
    """Test company name extraction and filename generation"""
    print("🧪 Testing company name extraction...")
    
    test_cases = [
        ("Software Developer at Innovation Labs", "We are Innovation Labs...", "Innovation_Labs"),
        ("Python Developer - Silverchair Silverchair", "About Silverchair...", "Silverchair"),
        ("Engineer @ Tech Corp Tech Corp", "Tech Corp is leading...", "Tech_Corp"), 
        ("Developer | Microsoft Microsoft Inc", "Microsoft Inc develops...", "Microsoft_Inc")
    ]
    
    for job_title, job_desc, expected in test_cases:
        company_name = extract_company_name(job_title, job_desc)
        print(f"Input: '{job_title}' → Extracted: '{company_name}'")
        
        # Test filename generation
        sample_personal_info = {"first_name": "John", "last_name": "Doe"}
        
        # Simulate the filename generation logic
        first_name = sample_personal_info.get('first_name', 'Resume').replace(' ', '')
        last_name = sample_personal_info.get('last_name', '').replace(' ', '')
        
        # Clean company name
        company_clean = re.sub(r'[^\w\s-]', '', company_name)
        company_clean = re.sub(r'\s+', ' ', company_clean).strip()
        
        # Remove duplicates
        company_words = company_clean.split(' ')
        unique_words = []
        seen_words = set()
        for word in company_words:
            word_lower = word.lower()
            if word and word_lower not in seen_words:
                unique_words.append(word)
                seen_words.add(word_lower)
        
        company_clean = '_'.join(unique_words)
        filename = f"{first_name}_{last_name}_{company_clean}.docx"
        
        print(f"  → Filename: {filename}")
        print(f"  → Expected pattern: John_Doe_{expected}.docx")
        print()

if __name__ == "__main__":
    print("Running tests...")
    test_company_name_extraction()
    print("\n" + "="*50 + "\n")
    test_resume_generation() 